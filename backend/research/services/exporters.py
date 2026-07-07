import re
from io import BytesIO


def report_filename(report, extension):
    return f"report-{report.id}.{extension}"


def render_report_pdf(report):
    try:
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_LEFT
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    except ImportError as exc:
        raise RuntimeError("后端未安装 reportlab，无法导出 PDF。请执行 pip install reportlab。") from exc

    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=18 * mm,
        leftMargin=18 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
        title=report.title,
    )

    base = getSampleStyleSheet()
    styles = {
        "title": ParagraphStyle(
            "ChineseTitle",
            parent=base["Title"],
            fontName="STSong-Light",
            fontSize=20,
            leading=28,
            textColor=colors.HexColor("#07183d"),
            spaceAfter=14,
            alignment=TA_LEFT,
        ),
        "heading1": ParagraphStyle(
            "ChineseHeading1",
            parent=base["Heading1"],
            fontName="STSong-Light",
            fontSize=15,
            leading=22,
            textColor=colors.HexColor("#12336f"),
            spaceBefore=12,
            spaceAfter=8,
        ),
        "heading2": ParagraphStyle(
            "ChineseHeading2",
            parent=base["Heading2"],
            fontName="STSong-Light",
            fontSize=13,
            leading=20,
            textColor=colors.HexColor("#1d4a92"),
            spaceBefore=10,
            spaceAfter=6,
        ),
        "body": ParagraphStyle(
            "ChineseBody",
            parent=base["BodyText"],
            fontName="STSong-Light",
            fontSize=10.5,
            leading=18,
            firstLineIndent=0,
            spaceAfter=7,
        ),
        "note": ParagraphStyle(
            "ChineseNote",
            parent=base["BodyText"],
            fontName="STSong-Light",
            fontSize=9.5,
            leading=16,
            textColor=colors.HexColor("#52627a"),
            backColor=colors.HexColor("#f5f8ff"),
            borderColor=colors.HexColor("#d7e3f7"),
            borderWidth=0.4,
            borderPadding=6,
            spaceAfter=10,
        ),
    }

    story = []
    for kind, text in _markdown_blocks(report.content_md):
        if not text:
            continue
        if kind == "title":
            story.append(Paragraph(_escape(text), styles["title"]))
        elif kind == "heading1":
            story.append(Paragraph(_escape(text), styles["heading1"]))
        elif kind == "heading2":
            story.append(Paragraph(_escape(text), styles["heading2"]))
        elif kind == "quote":
            story.append(Paragraph(_escape(text), styles["note"]))
        else:
            story.append(Paragraph(_escape(text), styles["body"]))
        story.append(Spacer(1, 2))

    doc.build(story)
    return buffer.getvalue()


def render_report_docx(report):
    try:
        from docx import Document
        from docx.oxml.ns import qn
        from docx.shared import Pt, RGBColor
    except ImportError as exc:
        raise RuntimeError("后端未安装 python-docx，无法导出 Word。请执行 pip install python-docx。") from exc

    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(10.5)

    for kind, text in _markdown_blocks(report.content_md):
        if not text:
            continue
        if kind == "title":
            paragraph = document.add_heading(text, level=0)
        elif kind == "heading1":
            paragraph = document.add_heading(text, level=1)
        elif kind == "heading2":
            paragraph = document.add_heading(text, level=2)
        else:
            paragraph = document.add_paragraph(text)
        for run in paragraph.runs:
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            if kind == "quote":
                run.font.color.rgb = RGBColor(82, 98, 122)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _markdown_blocks(markdown):
    for line in str(markdown or "").splitlines():
        text = line.strip()
        if not text:
            continue
        if text.startswith("# "):
            yield "title", text[2:].strip()
        elif text.startswith("## "):
            yield "heading1", text[3:].strip()
        elif text.startswith("### "):
            yield "heading2", text[4:].strip()
        elif text.startswith(">"):
            yield "quote", text.lstrip(">").strip()
        elif re.match(r"^[-*]\s+", text):
            yield "body", "• " + re.sub(r"^[-*]\s+", "", text)
        else:
            yield "body", _strip_markdown_inline(text)


def _strip_markdown_inline(text):
    value = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    value = re.sub(r"`([^`]+)`", r"\1", value)
    return value


def _escape(text):
    return (
        _strip_markdown_inline(text)
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )
